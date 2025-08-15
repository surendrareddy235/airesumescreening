import os
from typing import Dict, Any
import PyPDF2
from docx import Document
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file using PyPDF2"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        return text.strip()
    
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        raise

def extract_text_from_file(file_path: str) -> str:
    """Extract text from file based on extension"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_extension = Path(file_path).suffix.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def chunk_text(text: str, max_chunk_size: int = 512) -> list:
    """Split text into chunks for better processing"""
    words = text.split()
    chunks = []
    current_chunk = []
    current_size = 0
    
    for word in words:
        word_size = len(word) + 1  # +1 for space
        
        if current_size + word_size > max_chunk_size and current_chunk:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_size = word_size
        else:
            current_chunk.append(word)
            current_size += word_size
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

def preprocess_resume_text(text: str) -> Dict[str, Any]:
    """Preprocess resume text and extract basic information"""
    text_lower = text.lower()
    lines = text.split('\n')
    
    # Basic preprocessing results
    result = {
        'raw_text': text,
        'clean_text': ' '.join(text.split()),
        'word_count': len(text.split()),
        'has_email': '@' in text and '.' in text,
        'has_phone': any(char.isdigit() for char in text),
        'lines': [line.strip() for line in lines if line.strip()],
        'sections': []
    }
    
    # Identify potential sections
    section_keywords = [
        'experience', 'education', 'skills', 'projects', 
        'certifications', 'summary', 'objective', 'achievements'
    ]
    
    for line in result['lines']:
        line_lower = line.lower()
        for keyword in section_keywords:
            if keyword in line_lower and len(line.split()) <= 3:
                result['sections'].append({
                    'type': keyword,
                    'text': line
                })
                break
    
    return result
